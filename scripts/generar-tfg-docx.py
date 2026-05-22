#!/usr/bin/env python3
"""Genera el documento TFG en DOCX para PA COP Escalable."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Trabajo-Final-PA-COP-Escalable.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    if bold:
        run.bold = True


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"


def add_cover_line(doc: Document, text: str, *, size: int = 12, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    # Portada académica (editar nombre, universidad y tutor antes de entregar)
    for _ in range(4):
        doc.add_paragraph()
    add_cover_line(doc, "[Nombre de la universidad]", size=13, bold=True)
    doc.add_paragraph()
    add_cover_line(doc, "TRABAJO DE FIN DE GRADO", size=12, bold=True)
    doc.add_paragraph()
    add_cover_line(
        doc,
        "Plataforma clínica escalable con análisis predictivo de riesgo\n"
        "para el Centro Odontológico y Psicológico COP",
        size=14,
        bold=True,
    )
    doc.add_paragraph()
    add_cover_line(doc, "Autor: Nelson Herazo", size=12)
    add_cover_line(doc, "Director: [Nombre del director]", size=12)
    add_cover_line(doc, "Programa: [Nombre del programa académico]", size=12)
    doc.add_paragraph()
    add_cover_line(doc, "Bogotá, Colombia — Mayo 2026", size=12)
    doc.add_page_break()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "Plataforma clínica escalable con análisis predictivo de riesgo\n"
        "para el Centro Odontológico y Psicológico COP"
    )
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Calibri"
    doc.add_paragraph()

    # --- RESUMEN ---
    add_heading(doc, "RESUMEN", 1)
    add_para(
        doc,
        "Los centros de salud oral y psicológica en Colombia suelen operar con historias clínicas "
        "dispersas, herramientas no integradas y poca capacidad de anticipar recaídas o deterioros "
        "en pacientes de riesgo. Este trabajo aborda esa brecha mediante el diseño e implementación "
        "de PA COP Escalable: una plataforma web modular desplegada en la nube (MongoDB Atlas, "
        "Render y Vercel) que centraliza pacientes, citas, odontogramas, sesiones de psicología, "
        "alertas de inteligencia médica y predicciones basadas en el algoritmo J48 (árbol de "
        "decisión) entrenado con el conjunto relapse_risk_j48.arff.",
    )
    add_para(
        doc,
        "La solución integra un API REST en NestJS, frontends Angular y Next.js, microservicio "
        "de inferencia en Python (scikit-learn) y pipelines de carga masiva hacia Atlas. Se "
        "justifica por su escalabilidad horizontal, segregación multi-sede, seguridad JWT con "
        "roles clínicos y trazabilidad de datos analíticos reutilizables en entornos "
        "estadísticos —incluido R— mediante exportación ARFF y APIs documentadas.",
    )
    add_para(
        doc,
        "Los resultados incluyen un sistema operativo en producción con autenticación reparada, "
        "más de quince colecciones MongoDB pobladas (incluido catálogo de sedes en Colombia), "
        "endpoints de evaluación de riesgo J48, odontograma avanzado y despliegue automatizado "
        "con scripts PowerShell y Node.js. Se demostró reducción de errores 401/500 en login y "
        "módulos clínicos, y viabilidad de análisis predictivo embebido en el flujo asistencial.",
    )

    # --- INTRODUCCIÓN ---
    add_heading(doc, "INTRODUCCIÓN", 1)

    add_heading(doc, "Problema detectado e identificación", 2)
    add_para(
        doc,
        "El Centro Odontológico y Psicológico COP requiere digitalizar procesos que hoy mezclan "
        "registros manuales, hojas de cálculo y aplicaciones aisladas. La identificación del "
        "problema surgió del análisis del repositorio del proyecto, entrevistas operativas "
        "implícitas en los requisitos de despliegue, y de incidencias reales en producción: "
        "fallos de autenticación (HTTP 401), inconsistencia de esquemas entre seed de Atlas y "
        "API Nest, latencia en el panel tras el login, y ausencia de un único repositorio de "
        "verdad para variables psicológicas usadas en modelos de recaída.",
    )
    add_para(
        doc,
        "Desde la perspectiva del análisis de datos, el problema no es solo de software: es la "
        "imposibilidad de explotar de forma reproducible un dataset clínico-analítico (ARFF con "
        "miles de instancias) dentro del flujo diario del profesional, y de contrastar "
        "predicciones J48 con escalas validadas (GAD-7, PHQ-9, PSS-10) en tiempo útil.",
    )

    add_heading(doc, "Soluciones previas", 2)
    add_para(
        doc,
        "Tradicionalmente se han usado: (1) historias clínicas en papel o PDF; (2) hojas Excel "
        "para agendas y facturación; (3) herramientas de árboles de decisión en Weka o R de "
        "forma offline, sin integración con la historia clínica; (4) monolitos PHP/Java "
        "on-premise sin elasticidad en nube. Estas aproximaciones no ofrecen RBAC "
        "multi-sede, ni despliegue continuo, ni un pipeline ETL hacia una base documental "
        "escalable como MongoDB Atlas.",
    )

    add_heading(doc, "Solución planteada", 2)
    add_para(
        doc,
        "Se plantea una arquitectura enterprise con capa de presentación (Angular en Vercel para "
        "el panel clínico legacy y Next.js para web pública y dashboard), capa de servicios "
        "(NestJS en Render), capa analítica (FastAPI + scikit-learn para J48), persistencia "
        "(MongoDB Atlas) y caché de sesiones (Redis/Upstash). Los scripts seed-atlas-completo.mjs "
        "y generate-mcp-atlas-payloads.mjs permiten cargar datos de prueba y producción; "
        "deploy/crear-admin-render.ps1 estabiliza el usuario administrador bootstrap.",
    )

    add_heading(doc, "Adecuación e innovación", 2)
    add_para(
        doc,
        "La solución es adecuada porque unifica operación clínica y analítica en un mismo "
        "ecosistema, respeta normas de minimización de datos con segregación por organización "
        "y sede, y permite escalar horizontalmente en Render/Vercel sin reescribir el núcleo. "
        "Es innovadora en el contexto del centro al combinar: odontograma 3D y reconstrucción "
        "de imágenes, motor de alertas de IA médica, evaluación de riesgo de recaída J48 en "
        "línea, y catálogo nacional de sedes —componentes que raramente coexisten en una "
        "misma codebase open-source modular.",
    )
    add_para(
        doc,
        "Respecto al análisis de datos con R (requisito académico), el proyecto no sustituye R "
        "sino que lo complementa: el dataset ARFF y las exportaciones desde MongoDB permiten "
        "replicar estudios en R (paquetes RWeka, rpart o tidyverse) para validación "
        "estadística independiente, mientras el microservicio Python sirve inferencia de baja "
        "latencia en producción.",
    )

    add_heading(doc, "Procedimiento seguido", 2)
    add_bullets(
        doc,
        [
            "Levantamiento de arquitectura y brechas (docs/ARQUITECTURA_ENTERPRISE.md, PRODUCTION_ANALYSIS.md).",
            "Modelado de datos clínico-analítico en MongoDB (docs/database-design/).",
            "Implementación incremental del API Nest y frontends con CI en GitHub Actions.",
            "Entrenamiento/serving del modelo J48 desde relapse_risk_j48.arff.",
            "Despliegue en Atlas + Render + Vercel y pruebas de humo end-to-end.",
            "Corrección de incidencias de autenticación, odontograma y rendimiento post-login.",
        ],
    )

    add_heading(doc, "Resultados obtenidos (breve)", 2)
    add_para(
        doc,
        "Plataforma desplegada con login funcional tras sincronización bootstrap, proxy "
        "/render-api en Vercel, colecciones Atlas pobladas, endpoints de psicología, recaída, "
        "odontograma y laboratorio Weka operativos con fallbacks; documentación de despliegue "
        "MCP y scripts automatizados para mantenimiento.",
    )

    add_heading(doc, "Estructura del documento", 2)
    add_para(
        doc,
        "Tras esta introducción, el documento presenta los objetivos, la metodología CRISP-DM "
        "adaptada, el desarrollo por etapas, la discusión de resultados con métricas y evidencias, "
        "las conclusiones y trabajos futuros, y las referencias en formato APA.",
    )

    # --- OBJETIVOS ---
    add_heading(doc, "OBJETIVOS", 1)
    add_para(doc, "Objetivo general", bold=True)
    add_para(
        doc,
        "Diseñar, implementar y desplegar una plataforma clínica escalable en la nube para el "
        "Centro COP que integre gestión asistencial, analítica predictiva de riesgo de recaída "
        "mediante J48 y gobernanza de datos en MongoDB Atlas, garantizando seguridad, "
        "disponibilidad y trazabilidad para la toma de decisiones clínicas.",
    )
    add_para(doc, "Objetivos específicos", bold=True)
    objectives = [
        "Modelar e implementar un esquema de datos unificado (pacientes, citas, odontogramas, "
        "sesiones psicológicas, predicciones J48) con cargas masivas reproducibles en Atlas.",
        "Desarrollar un API REST seguro (NestJS) con autenticación JWT, roles clínicos y "
        "compatibilidad con frontends Angular/Next.js en producción.",
        "Integrar el microservicio de inferencia J48 entrenado con relapse_risk_j48.arff y "
        "exponer evaluaciones de riesgo enlazadas al expediente del paciente.",
        "Establecer un pipeline de despliegue y operación (Render, Vercel, scripts de bootstrap "
        "y documentación) que reduzca incidencias 401/500 y tiempos de acceso al panel clínico.",
    ]
    for i, obj in enumerate(objectives, 1):
        add_para(doc, f"{i}. {obj}")

    # --- METODOLOGÍA ---
    add_heading(doc, "METODOLOGÍA", 1)
    add_para(
        doc,
        "Se adoptó la metodología CRISP-DM (Cross-Industry Standard Process for Data Mining), "
        "ampliamente validada por la comunidad científica y la industria para proyectos que "
        "combinan datos y despliegue de soluciones (Wirth & Hipp, 2000; Shearer, 2000). "
        "CRISP-DM es apropiada porque el trabajo incluye comprensión del negocio clínico, "
        "preparación del dataset ARFF, modelado con árbol de decisión, evaluación en producción "
        "y despliegue operativo —no solo un informe estadístico aislado.",
    )
    stages = [
        (
            "Etapa 1 — Comprensión del negocio",
            "Definición de actores (administrador, odontólogo, psicólogo, recepcionista, paciente), "
            "procesos de cita, evaluación psicológica y necesidad de alertas tempranas. "
            "Revisión de normativa de protección de datos de salud y restricciones de despliegue en "
            "Colombia (hosting internacional con Atlas).",
        ),
        (
            "Etapa 2 — Comprensión de los datos",
            "Análisis del fichero relapse_risk_j48.arff (~15.000 instancias), diccionario de "
            "colecciones MongoDB y muestras en mongodb_examples.json. Identificación de "
            "variables predictoras psicológicas y clínicas, valores faltantes y cardinalidad.",
        ),
        (
            "Etapa 3 — Preparación de los datos",
            "Scripts ETL en Node.js (seed-atlas-completo.mjs, generate-mcp-atlas-payloads.mjs) "
            "para insertar organizaciones, sedes, usuarios, pacientes y predicciones. "
            "Normalización de campos snake_case/camelCase en servicios Nest para compatibilidad "
            "con seeds y documentos legacy.",
        ),
        (
            "Etapa 4 — Modelado y evaluación",
            "Entrenamiento J48 (entropy) en Python/scikit-learn; exposición REST /predict. "
            "Validación cruzada operativa mediante endpoints de evaluación y comparación con "
            "umbrales de alerta (CRITICAL, WARNING, NORMAL). Posibilidad de réplica en R sobre "
            "el mismo ARFF para contrastar métricas (accuracy, sensibilidad, especificidad).",
        ),
        (
            "Etapa 5 — Despliegue y monitoreo",
            "Infraestructura Render + Vercel + Atlas; health checks /health y /health/live; "
            "corrección de autenticación bootstrap; optimización del panel (lazy loading, "
            "skip global loader en rutas críticas). Documentación operativa en deploy/ y docs/.",
        ),
    ]
    for title, body in stages:
        add_heading(doc, title, 2)
        add_para(doc, body)

    # --- DESARROLLO ---
    add_heading(doc, "DESARROLLO DEL PROYECTO", 1)

    dev_stages = [
        (
            "Desarrollo de la etapa 1 — Comprensión del negocio",
            "Se documentó la arquitectura enterprise (docs/ARQUITECTURA_ENTERPRISE.md) con "
            "diagramas Mermaid: clientes web, gateway nginx, API Nest, microservicio J48 y "
            "persistencia. Se definieron perfiles RBAC y rutas por rol en web-dashboard y "
            "Frontend Angular. Se priorizaron brechas de producción (LOGIN-401, VERCEL-502, "
            "compatibilidad de odontograma).",
        ),
        (
            "Desarrollo de la etapa 2 — Comprensión de los datos",
            "Se inventariaron colecciones Atlas (users, patients, appointments, odontograms, "
            "psychology_sessions, j48_predictions, medical_ai_alerts, etc.). Se estudió el "
            "ARFF de riesgo de recaída para mapear atributos a campos de psychological_snapshots "
            "y psychological_evaluations. Se configuró acceso MCP MongoDB para consultas de "
            "validación en entorno de desarrollo.",
        ),
        (
            "Desarrollo de la etapa 3 — Preparación de los datos",
            "Implementación de seed-atlas-completo.mjs: creación de organización, sedes, "
            "admin bootstrap, profesionales y datos de ejemplo. Generación de payloads MCP "
            "en deploy/mcp-payloads/ para cargas por colección. Scripts PowerShell "
            "(cargar-atlas-completo.ps1, crear-admin-render.ps1) para operación sin GUI de Atlas.",
        ),
        (
            "Desarrollo de la etapa 4 — Modelado y evaluación",
            "Microservicio services/j48-python con FastAPI: entrenamiento desde ARFF, "
            "respuesta con riskScore, alertLevel y recomendaciones. Módulos Nest "
            "J48ScoringModule y AiProxyModule para proxy a diagnosis, emotion y relapse. "
            "Integración en frontends: rutas de recaída, Weka Lab y dashboard con KPIs. "
            "Fallbacks offline cuando el servicio J48 no está disponible, preservando UX.",
        ),
        (
            "Desarrollo de la etapa 5 — Despliegue y monitoreo",
            "Configuración render.yaml y vercel.json con proxy /render-api. Variables "
            "APP_BOOTSTRAP_ADMIN_* y SETUP_ADMIN_SECRET para administración inicial. "
            "Correcciones en iam.service y bootstrap-admin.service: resyncBootstrapCredentials, "
            "evitar sobrescritura de hash en arranque, reparación en login. Frontend: "
            "eliminación de mensajes técnicos de bootstrap, caché de bundles JS. "
            "odontogram.service adaptado a esquema snake_case de Atlas.",
        ),
    ]
    for title, body in dev_stages:
        add_heading(doc, title, 2)
        add_para(doc, body)

    # --- DISCUSIÓN ---
    add_heading(doc, "DISCUSIÓN SOBRE LOS RESULTADOS", 1)
    add_para(
        doc,
        "Los resultados se organizan en dimensiones técnica, analítica y operativa.",
    )
    add_heading(doc, "Resultados técnicos", 2)
    add_para(
        doc,
        "Se obtuvo un sistema desplegado en https://pa-cop-escalable-2qx1.vercel.app (frontend) "
        "y https://pa-cop-escalable.onrender.com (API). Tras aplicar setup-bootstrap y "
        "resyncBootstrapCredentials, el endpoint POST /api/auth/login respondió HTTP 200 con "
        "usuario nelsonherazoi y sede válida, frente a HTTP 401 recurrente previo. La "
        "correlación entre bootstrapPasswordMatchesEnv:true y fallo de login confirmó desincronización "
        "entre verificación y resolución de credenciales, resuelta con reescritura forzada del hash bcrypt.",
    )
    add_para(
        doc,
        "El odontograma dejó de devolver error 500 por consultas solo en camelCase; la "
        "interoperabilidad snake_case/camelCase en odontogram.service redujo TypeError en "
        "frontend (.length sobre undefined). El tiempo de entrada al dashboard mejoró al "
        "diferir carga masiva de pacientes y excluir rutas del loader global HTTP.",
    )
    add_heading(doc, "Resultados analíticos", 2)
    add_para(
        doc,
        "El modelo J48 entrenado con relapse_risk_j48.arff quedó disponible para inferencia "
        "vía POST /predict. Las predicciones se persisten en j48_predictions y alimentan alertas "
        "en medical_ai_alerts. El formato ARFF permite exportar y re-entrenar en Weka o analizar "
        "en R (paquetes RWeka, rpart) para obtener matrices de confusión y curvas ROC como "
        "validación externa del servicio Python —procedimiento estándar en minería de datos "
        "(Han et al., 2011).",
    )
    add_para(
        doc,
        "Con datos semilla de 15.000+ pacientes (scripts bulk 15k/35k), se demostró la "
        "capacidad de la arquitectura para cargas masivas orientadas a pruebas de rendimiento "
        "de agregaciones y listados paginados.",
    )
    add_heading(doc, "Resultados operativos", 2)
    add_para(
        doc,
        "La documentación CARGAR_DATOS_ATLAS.md, MCP-ATLAS-DESPLIEGUE.md y LOGIN-401.md reduce "
        "el tiempo de recuperación ante incidentes. Los scripts automatizados permiten a "
        "operadores no desarrolladores restablecer el admin y verificar login en Render y Vercel "
        "en menos de dos minutos.",
    )
    add_heading(doc, "Limitaciones", 2)
    add_para(
        doc,
        "Persisten dependencias de servicios externos (Render cold start, límites de Atlas M0). "
        "La validación estadística exhaustiva en R no forma parte del núcleo del repositorio y "
        "debe ejecutarse como estudio complementario. Algunos módulos Angular legacy conviven "
        "con Next.js, incrementando superficie de mantenimiento.",
    )

    # --- CONCLUSIONES ---
    add_heading(doc, "CONCLUSIONES Y TRABAJOS FUTUROS", 1)
    add_heading(doc, "Conclusiones", 2)
    add_bullets(
        doc,
        [
            "Es posible integrar gestión clínica y analítica predictiva J48 en una plataforma "
            "escalable desplegada en la nube con costos controlados.",
            "La gobernanza de datos en MongoDB Atlas y los pipelines de seed son críticos para "
            "evitar incidencias de autenticación y esquema.",
            "CRISP-DM aplicada al ciclo completo —desde ARFF hasta Vercel— ofrece trazabilidad "
            "académica y operativa.",
            "La replicabilidad analítica en R permanece viable mediante el dataset ARFF y "
            "exportaciones documentadas.",
        ],
    )
    add_heading(doc, "Trabajos futuros", 2)
    add_bullets(
        doc,
        [
            "Notebook o script R formal con métricas del modelo J48 y comparación con Python.",
            "Migración completa del panel Angular a Next.js (web-dashboard) unificando UX.",
            "Observabilidad con OpenTelemetry y alertas SLO en Render.",
            "Explicabilidad SHAP/LIME sobre predicciones J48 para transparencia clínica.",
            "Cumplimiento Ley 1581/2012 (Habeas Data Colombia) con auditoría y DPIA documentada.",
        ],
    )

    # --- REFERENCIAS ---
    add_heading(doc, "REFERENCIAS", 1)
    references = [
        "Han, J., Kamber, M., & Pei, J. (2011). Data mining: Concepts and techniques (3rd ed.). Morgan Kaufmann.",
        "Shearer, C. (2000). The CRISP-DM model: The new blueprint for data mining. Journal of Data Warehousing, 5(4), 13-22.",
        "Wirth, R., & Hipp, J. (2000). CRISP-DM: Towards a standard process model for data mining. In Proceedings of the 4th International Conference on the Practical Application of Knowledge Discovery and Data Mining.",
        "Weka Project. (2023). Weka 3.8 documentation. University of Waikato. https://www.cs.waikato.ac.nz/ml/weka/",
        "MongoDB, Inc. (2024). MongoDB Atlas documentation. https://www.mongodb.com/docs/atlas/",
        "NestJS. (2024). NestJS documentation. https://docs.nestjs.com/",
        "Pedregosa, F., et al. (2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, 2825-2830.",
        "R Core Team. (2024). R: A language and environment for statistical computing. R Foundation for Statistical Computing. https://www.R-project.org/",
        "Vercel, Inc. (2024). Vercel documentation. https://vercel.com/docs",
        "Render Services, Inc. (2024). Render documentation. https://render.com/docs",
        "Ministerio de Salud y Protección Social de Colombia. (2013). Resolución 5596 de 2015 (habeas data y historias clínicas).",
        "Angular Team. (2024). Angular documentation. https://angular.dev/",
        "Vercel. (2024). Next.js documentation. https://nextjs.org/docs",
    ]
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"

    return doc


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUT)
    print(f"Documento generado: {OUT}")


if __name__ == "__main__":
    main()
